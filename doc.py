import docx


doc = docx.Document('master_file.docx')
def doc_open_save(DISC, special, spec, kafedra, uroven, compets, pod_competions, year, author, nach_kaf):
    doc_tab_DISC(DISC)
    doc_str_DISC(DISC)
    doc_tab_kafedra(kafedra)
    doc_str_kafedra(kafedra)
    doc_str_spec(spec)
    doc_str_special(special)
    doc_str_uroven(uroven)
    competentions(compets, pod_competions)
    tab_competentions(compets, pod_competions)
    tab_year(year)
    str_year(year)
    tab_author(author)
    str_author(author)
    tab_nach_kaf(nach_kaf)
    str_nach_kaf(nach_kaf)
    doc.save("test.docx")


def doc_tab_DISC(DISC):
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.find("DISC") > -1:
                        for run in para.runs:
                            if run.text == "DISC":
                                run.text = str(DISC)
                                
                        

def doc_str_DISC(DISC):
    for para in doc.paragraphs:
        if para.text.find("DISC") > -1:
            for run in para.runs:
                if run.text.find("DISC") > - 1:
                    run.text = run.text.replace("DISC", DISC)


def doc_tab_kafedra(kafedra):
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.find("kafedra") > -1:
                        for run in para.runs:
                            if run.text.find("kafedra") > -1:
                                run.text = "Кафедра «" + kafedra +"»"
                                

def doc_str_kafedra(kafedra):
    for para in doc.paragraphs:
        if para.text.find("kafedra") > -1:
            para.text = para.text.replace("kafedra", kafedra)                               

def doc_str_spec(spec):
    for para in doc.paragraphs:
        if para.text.lower().find("spek") > -1:
            para.text = para.text.replace("spek", spec)
                              

def doc_str_special(special):
    for para in doc.paragraphs:
        if para.text.find("special") > -1:
            for run in para.runs:
                if run.text.find("special") > -1:
                    run.text = run.text.replace("special", special)

def doc_str_uroven(uroven):
    for para in doc.paragraphs:
        if para.text.lower().find("uroven") > -1:
            para.text = para.text.replace("uroven", uroven)

def competentions(compets, pod_competions):
    comps = ""
    for para in doc.paragraphs:
        if para.text.find("Competentions") > -1:
            for com in compets:
                comps = comps + com["code"] + " " + com["description"] + '\n'
                for podcom in pod_competions:
                    if com['id'] == podcom['id_of_comp']:
                        comps = comps + podcom["code_of_pod_comp"] + " " + podcom['descript'] + '\n'
            para.text = comps


def tab_competentions(compets, pod_competions):
    for tab in doc.tables:
        for row in tab.rows:
            if row.cells[0].text.lower().find("code") > -1:
                for i in range(len(compets) - 1):
                    tab.add_row()
                
    for tab in doc.tables:
        comp = compets[:]
        if len(tab.rows) > 1:
            if tab.rows[1].cells[0].text.lower().find("code") > -1:
                for i in range(len(tab.rows) - 1):
                    tab.rows[i + 1].cells[0].text = comp[0]["code"] + "code"
                    tab.rows[i + 1].cells[1].text = comp[0]["description"]
                    pod_com = ""
                    for pods in pod_competions:
                        if pods['id_of_comp'] == comp[0]['id']:
                            pod_com = pod_com + pods["code_of_pod_comp"] + " " + pods['descript'] + '\n'
                    tab.rows[i + 1].cells[2].text = pod_com
                    comp.remove(comp[0])
        for row in tab.rows:
            if row.cells[0].text.find('code') > -1:
                row.cells[0].text = row.cells[0].text.replace("code", "")

def tab_year(year):
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                       if para.text.find("year") > -1:
                                para.text = para.text.replace("year", year)

def str_year(year):
    for para in doc.paragraphs:
        if para.text.find("year") > -1:
            para.text =  para.text.replace("year", year)

def tab_author(author):
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                       if para.text.find("author") > -1:
                                para.text = para.text.replace("author", author)

def str_author(author):
    for para in doc.paragraphs:
        if para.text.find("author") > -1:
            para.text =  para.text.replace("author", author)

def tab_nach_kaf(nach_kaf):
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                       if para.text.find("nach_kaf") > -1:
                                para.text = para.text.replace("nach_kaf", nach_kaf)

def str_nach_kaf(nach_kaf):
    for para in doc.paragraphs:
        if para.text.find("nach_kaf") > -1:
            para.text =  para.text.replace("nach_kaf", nach_kaf)
